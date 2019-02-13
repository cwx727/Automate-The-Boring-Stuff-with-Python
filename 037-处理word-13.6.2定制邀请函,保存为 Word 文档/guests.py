import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = docx.Document()
with open('guests.txt') as obj_f:
	names = obj_f.readlines()

for name in names:
	p = doc.add_paragraph('')
	p.add_run('It would be a pleasure to have the company of').italic = True
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	p = doc.add_paragraph('')
	p.add_run(name.rstrip()).bold = True
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	p = doc.add_paragraph('')
	p.add_run('at 11010 Menary Lane on the Evening of').italic = True
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	p = doc.add_paragraph('')
	p.add_run('April 1st')
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	p = doc.add_paragraph('')
	p.add_run("at 7 o'clock").italic = True
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc.add_page_break()


doc.save('guests.docx')
